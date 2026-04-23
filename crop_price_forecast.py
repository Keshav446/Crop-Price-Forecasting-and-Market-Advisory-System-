"""
=============================================================================
  AI-Based Crop Price Forecasting & Market Advisory System
  ─────────────────────────────────────────────────────────
  Model   : Support Vector Regression (SVR – RBF kernel)
  Metrics : MAE · RMSE · R²  (Training + Testing)
  Output  : Console metrics, sorted + scatter graphs, Word report
  Data    : train_data_150k.csv  /  test_data_60k.csv
=============================================================================
"""

# ── Libraries ────────────────────────────────────────────────────────────────
import os, sys, math
import numpy as np
import pandas as pd
from sklearn.svm import SVR
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
import joblib
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
warnings.filterwarnings("ignore")

# ── Paths & Constants ────────────────────────────────────────────────────────
TRAIN_PATH       = "train_data_150k.csv"
TEST_PATH        = "test_data_60k.csv"
TARGET_COL       = "price"
MODEL_PATH       = "svr_model.pkl"
SCALER_PATH      = "scaler.pkl"
REPORT_PATH      = "Crop_Price_Report.docx"
TRAIN_SORTED_PNG = "train_sorted_plot.png"
TEST_SORTED_PNG  = "test_sorted_plot.png"
SCATTER_PNG      = "scatter_plot.png"
GRAPH_SAMPLES    = 1000   # for sorted line plots only


# ═══════════════════════════════════════════════════════════════════════════════
#  1. DATA LOADING
# ═══════════════════════════════════════════════════════════════════════════════

def load_data(train_path, test_path):
    """Load FULL CSVs — no sampling, no row limits."""
    for p in (train_path, test_path):
        if not os.path.isfile(p):
            sys.exit(f"[ERROR] File not found: '{p}'")

    train_df = pd.read_csv(train_path)
    test_df  = pd.read_csv(test_path)

    print("Loaded real dataset successfully")
    print(f"Actual training rows loaded: {len(train_df)}")
    print(f"Actual testing rows loaded:  {len(test_df)}")
    print(f"\n[INFO] Training data shape : {train_df.shape}")
    print(f"[INFO] Testing  data shape : {test_df.shape}")

    print("\n── First 5 rows of TRAINING data (preview — ALL rows are used) ──")
    print(train_df.head().to_string(index=False))
    print("\n── First 5 rows of TESTING data (preview — ALL rows are used) ──")
    print(test_df.head().to_string(index=False))

    return train_df, test_df


# ═══════════════════════════════════════════════════════════════════════════════
#  2. PREPROCESSING
# ═══════════════════════════════════════════════════════════════════════════════

def handle_missing_values(df, label="data"):
    """Numerical → median, categorical → mode."""
    missing = df.isnull().sum().sum()
    for col in df.columns:
        if df[col].isnull().sum() == 0:
            continue
        if df[col].dtype in ("float64", "int64", "float32", "int32"):
            df[col].fillna(df[col].median(), inplace=True)
        else:
            df[col].fillna(df[col].mode()[0], inplace=True)
    print(f"[INFO] Missing values in {label}: {missing} → 0  (filled)")
    return df


def encode_categorical(train_df, test_df):
    """Label-encode categoricals — fit on TRAINING only."""
    encoders = {}
    cat_cols = train_df.select_dtypes(include=["object", "category"]).columns.tolist()
    for col in cat_cols:
        le = LabelEncoder()
        le.fit(train_df[col].astype(str))
        train_df[col] = le.transform(train_df[col].astype(str))
        test_df[col]  = le.transform(test_df[col].astype(str))
        encoders[col] = le
    if cat_cols:
        print(f"[INFO] Encoded categorical columns: {cat_cols}")
    return train_df, test_df, encoders


def scale_features(X_train, X_test):
    """StandardScaler — fit on train only."""
    scaler = StandardScaler()
    X_train_s = scaler.fit_transform(X_train)
    X_test_s  = scaler.transform(X_test)
    print("[INFO] Feature scaling applied (StandardScaler — fit on train only).")
    return X_train_s, X_test_s, scaler


def preprocess(train_df, test_df, target_col):
    """Full pipeline: missing → encode → split → scale."""
    train_df = handle_missing_values(train_df.copy(), "training set")
    test_df  = handle_missing_values(test_df.copy(),  "testing set")
    train_df, test_df, encoders = encode_categorical(train_df, test_df)

    feat_cols = [c for c in train_df.columns if c != target_col]
    X_train = train_df[feat_cols].values;  y_train = train_df[target_col].values
    X_test  = test_df[feat_cols].values;   y_test  = test_df[target_col].values

    X_train_s, X_test_s, scaler = scale_features(X_train, X_test)
    print(f"[INFO] Features used: {feat_cols}")
    return X_train_s, X_test_s, y_train, y_test, scaler, encoders, feat_cols


# ═══════════════════════════════════════════════════════════════════════════════
#  3. MODEL
# ═══════════════════════════════════════════════════════════════════════════════

SVR_PARAMS = {"kernel": "rbf", "C": 10.0, "gamma": "scale", "epsilon": 0.1}


def build_svr(**params):
    model = SVR(**params)
    print(f"\n[MODEL] SVR created  →  {params}")
    return model


def train_model(model, X, y):
    print("[INFO] Training on full dataset — this may take a moment...")
    model.fit(X, y)
    print("[INFO] ✅  Model training complete.\n")
    return model


# ═══════════════════════════════════════════════════════════════════════════════
#  4. METRICS
# ═══════════════════════════════════════════════════════════════════════════════

def compute_metrics(y_true, y_pred, label=""):
    mse  = mean_squared_error(y_true, y_pred)
    mae  = mean_absolute_error(y_true, y_pred)
    rmse = math.sqrt(mse)
    r2   = r2_score(y_true, y_pred)

    print("=" * 55)
    print(f"       {label} METRICS")
    print("=" * 55)
    print(f"  Mean Absolute Error (MAE)  : {mae:>12.4f}")
    print(f"  Root Mean Sq Error  (RMSE) : {rmse:>12.4f}")
    print(f"  R² Score                   : {r2:>12.4f}")
    print("=" * 55 + "\n")

    return {"mae": mae, "mse": mse, "rmse": rmse, "r2": r2}


# ═══════════════════════════════════════════════════════════════════════════════
#  5. RECOMMENDATION LOGIC  (FIXED — per-sample actual vs predicted)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_recommendations(y_actual, y_pred,
                             hold_thresh=0.10, wait_thresh=0.02):
    """
    Per-sample comparison:
      change = (predicted - actual) / actual
      ≥ +10 %   → HOLD
      +2 % – 10 % → WAIT
      < +2 %    → SELL
    """
    recs = []
    for actual, pred in zip(y_actual, y_pred):
        change = (pred - actual) / actual
        if change >= hold_thresh:
            recs.append("HOLD")
        elif change >= wait_thresh:
            recs.append("WAIT")
        else:
            recs.append("SELL")
    return recs


# ═══════════════════════════════════════════════════════════════════════════════
#  6. CONSOLE DISPLAY
# ═══════════════════════════════════════════════════════════════════════════════

def display_predictions(y_true, y_pred, recs, n=20):
    print(f"\n── Sample Predictions (first {n} of {len(y_true)}) ──\n")
    print("=" * 70)
    print(f"  {'#':>4}   {'Actual':>12}   {'Predicted':>12}   {'Advisory':>10}")
    print("-" * 70)
    for i in range(min(n, len(y_true))):
        print(f"  {i+1:>4}   {y_true[i]:>12.2f}   {y_pred[i]:>12.2f}   {recs[i]:>10}")
    print("=" * 70)


# ═══════════════════════════════════════════════════════════════════════════════
#  7. GRAPH VISUALISATION  (FIXED — sorted line plots + scatter)
# ═══════════════════════════════════════════════════════════════════════════════

def plot_sorted(y_true, y_pred, title, save_path, n=GRAPH_SAMPLES):
    """
    Sort by actual price before plotting → smooth, interpretable curves.
    Only first n samples are plotted for visual clarity.
    """
    sorted_idx   = np.argsort(y_true)
    y_true_sort  = y_true[sorted_idx]
    y_pred_sort  = y_pred[sorted_idx]

    n = min(n, len(y_true_sort))
    idx = np.arange(n)

    plt.figure(figsize=(14, 5))
    plt.plot(idx, y_true_sort[:n], color="#2196F3", linewidth=1.2,
             alpha=0.9, label="Actual Price")
    plt.plot(idx, y_pred_sort[:n], color="#FF5722", linewidth=1.2,
             alpha=0.9, label="Predicted Price")
    plt.xlabel("Sample Index (sorted by actual price)", fontsize=12)
    plt.ylabel("Price (₹)", fontsize=12)
    plt.title(title, fontsize=14, fontweight="bold")
    plt.legend(fontsize=11)
    plt.tight_layout()
    plt.savefig(save_path, dpi=150)
    plt.close()
    print(f"[INFO] 📊  Graph saved → {save_path}")


def plot_scatter(y_true, y_pred, save_path):
    """
    Scatter plot of actual vs predicted with a y = x diagonal reference line.
    """
    plt.figure(figsize=(8, 8))
    plt.scatter(y_true, y_pred, alpha=0.3, s=8, color="#1976D2", label="Predictions")

    # diagonal reference line (perfect prediction)
    lo = min(y_true.min(), y_pred.min())
    hi = max(y_true.max(), y_pred.max())
    plt.plot([lo, hi], [lo, hi], color="#E53935", linewidth=2,
             linestyle="--", label="Ideal (y = x)")

    plt.xlabel("Actual Price (₹)", fontsize=12)
    plt.ylabel("Predicted Price (₹)", fontsize=12)
    plt.title("Actual vs Predicted Scatter Plot", fontsize=14, fontweight="bold")
    plt.legend(fontsize=11)
    plt.tight_layout()
    plt.savefig(save_path, dpi=150)
    plt.close()
    print(f"[INFO] 📊  Graph saved → {save_path}")


# ═══════════════════════════════════════════════════════════════════════════════
#  8. WORD REPORT  (UPGRADED — all 3 graphs + explanations)
# ═══════════════════════════════════════════════════════════════════════════════

def _cell(cell, text, bold=False, center=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.font.size = Pt(10)
    run.bold = bold
    if center:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_metrics_section(doc, heading, metrics):
    doc.add_heading(heading, level=1)
    tm = doc.add_table(rows=1, cols=2, style="Light Shading Accent 1")
    _cell(tm.rows[0].cells[0], "Metric", bold=True, center=True)
    _cell(tm.rows[0].cells[1], "Value",  bold=True, center=True)
    for key, label in [("mae", "MAE"), ("rmse", "RMSE"), ("r2", "R² Score")]:
        r = tm.add_row()
        _cell(r.cells[0], label, center=True)
        _cell(r.cells[1], f"{metrics[key]:.4f}", center=True)


def generate_word_report(train_rows, test_rows, svr_params,
                         train_metrics, test_metrics,
                         y_true, y_pred, recs, report_path):
    doc = Document()

    # Title
    t = doc.add_heading("Crop Price Forecasting Report", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1 — Dataset Summary
    doc.add_heading("1. Dataset Summary", level=1)
    doc.add_paragraph(f"Training samples : {train_rows:,}")
    doc.add_paragraph(f"Testing  samples : {test_rows:,}")

    # 2 — Model Details
    doc.add_heading("2. Model Details", level=1)
    doc.add_paragraph("Algorithm: Support Vector Regression (SVR)")
    tp = doc.add_table(rows=1, cols=2, style="Light Shading Accent 1")
    _cell(tp.rows[0].cells[0], "Parameter", bold=True, center=True)
    _cell(tp.rows[0].cells[1], "Value",     bold=True, center=True)
    for k, v in svr_params.items():
        r = tp.add_row()
        _cell(r.cells[0], k,      center=True)
        _cell(r.cells[1], str(v), center=True)

    # 3 & 4 — Metrics
    _add_metrics_section(doc, "3. Training Metrics", train_metrics)
    _add_metrics_section(doc, "4. Testing Metrics",  test_metrics)

    # 5 — Visualisations (all 3 graphs)
    doc.add_heading("5. Visualisations", level=1)

    if os.path.isfile(TRAIN_SORTED_PNG):
        doc.add_paragraph("Training Data (Sorted): Actual vs Predicted")
        doc.add_picture(TRAIN_SORTED_PNG, width=Inches(6.0))
        doc.add_paragraph(
            "Sorted plots improve trend visibility by ordering samples by "
            "actual price, making it easy to see how well predictions track "
            "the true price curve."
        )

    if os.path.isfile(TEST_SORTED_PNG):
        doc.add_paragraph("Testing Data (Sorted): Actual vs Predicted")
        doc.add_picture(TEST_SORTED_PNG, width=Inches(6.0))
        doc.add_paragraph(
            "The sorted test plot shows prediction accuracy across the full "
            "price range, revealing any systematic bias at low or high prices."
        )

    if os.path.isfile(SCATTER_PNG):
        doc.add_paragraph("Actual vs Predicted Scatter Plot")
        doc.add_picture(SCATTER_PNG, width=Inches(5.5))
        doc.add_paragraph(
            "The scatter plot shows prediction accuracy — points closer to "
            "the diagonal y = x line indicate better predictions. Tight "
            "clustering around the line confirms strong model performance."
        )

    # 6 — Sample Predictions Table
    doc.add_heading("6. Sample Predictions (First 20 — Test Data)", level=1)
    n = min(20, len(y_true))
    pt = doc.add_table(rows=1, cols=3, style="Light Shading Accent 1")
    for i, h in enumerate(["Actual Price", "Predicted Price", "Recommendation"]):
        _cell(pt.rows[0].cells[i], h, bold=True, center=True)
    for i in range(n):
        r = pt.add_row()
        _cell(r.cells[0], f"{y_true[i]:.2f}", center=True)
        _cell(r.cells[1], f"{y_pred[i]:.2f}", center=True)
        _cell(r.cells[2], recs[i],             center=True)

    doc.save(report_path)
    print(f"[INFO] 📄  Word report saved → {report_path}")


# ═══════════════════════════════════════════════════════════════════════════════
#  9. SAVE / LOAD MODEL
# ═══════════════════════════════════════════════════════════════════════════════

def save_artifacts(model, scaler):
    joblib.dump(model,  MODEL_PATH)
    joblib.dump(scaler, SCALER_PATH)
    print(f"[INFO] Model  saved → {MODEL_PATH}")
    print(f"[INFO] Scaler saved → {SCALER_PATH}")


def load_artifacts():
    model  = joblib.load(MODEL_PATH)
    scaler = joblib.load(SCALER_PATH)
    print(f"[INFO] Model  loaded ← {MODEL_PATH}")
    print(f"[INFO] Scaler loaded ← {SCALER_PATH}")
    return model, scaler


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print("\n" + "═" * 55)
    print("  AI-Based Crop Price Forecasting & Market Advisory")
    print("═" * 55 + "\n")

    # 1 — Load
    train_df, test_df = load_data(TRAIN_PATH, TEST_PATH)

    # 2 — Preprocess
    (X_train, X_test, y_train, y_test,
     scaler, encoders, feat_names) = preprocess(train_df, test_df, TARGET_COL)

    # 3 — Build & Train
    model = build_svr(**SVR_PARAMS)
    model = train_model(model, X_train, y_train)

    # 4 — Predict on BOTH sets
    print("[INFO] Predicting on training data...")
    y_pred_train = model.predict(X_train)
    print("[INFO] Predicting on testing data...")
    y_pred_test  = model.predict(X_test)

    # 5 — Metrics (train + test)
    train_metrics = compute_metrics(y_train, y_pred_train, label="TRAINING")
    test_metrics  = compute_metrics(y_test,  y_pred_test,  label="TESTING")

    # 6 — Recommendations (FIXED: per-sample actual vs predicted)
    recs = generate_recommendations(y_test, y_pred_test)

    # 7 — Console display
    display_predictions(y_test, y_pred_test, recs, n=20)

    # 8 — Graphs (FIXED: sorted + scatter)
    plot_sorted(y_train, y_pred_train,
                "Training Data (Sorted): Actual vs Predicted",
                TRAIN_SORTED_PNG)
    plot_sorted(y_test, y_pred_test,
                "Testing Data (Sorted): Actual vs Predicted",
                TEST_SORTED_PNG)
    plot_scatter(y_test, y_pred_test, SCATTER_PNG)

    # 9 — Word report (with all 3 graphs)
    generate_word_report(
        train_rows=len(y_train), test_rows=len(y_test),
        svr_params=SVR_PARAMS,
        train_metrics=train_metrics, test_metrics=test_metrics,
        y_true=y_test, y_pred=y_pred_test, recs=recs,
        report_path=REPORT_PATH,
    )

    # 10 — Save & verify
    save_artifacts(model, scaler)

    print("\n── Reloading saved model for verification ──")
    loaded_model, _ = load_artifacts()
    y_check = loaded_model.predict(X_test)
    if np.allclose(y_pred_test, y_check):
        print("[INFO] ✅  Loaded model produces identical predictions.\n")
    else:
        print("[WARNING] ⚠️  Loaded model predictions differ!\n")


if __name__ == "__main__":
    main()
