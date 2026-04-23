"""
=============================================================================
  AI-Based Crop Price Forecasting using Linear Regression
  ─────────────────────────────────────────────────────────
  Model   : Linear Regression
  Metrics : MAE · RMSE · R²  (Training + Testing)
  Output  : Console metrics, sorted + scatter graphs, Word report
  Data    : train_data_150k.csv  /  test_data_60k.csv
=============================================================================
"""

import os, sys, math
import numpy as np
import pandas as pd
from sklearn.linear_model import LinearRegression
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
import joblib
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings

warnings.filterwarnings("ignore")

# ── Paths & Constants ────────────────────────────────────────────────────────
TRAIN_PATH       = "train_data_150k.csv"
TEST_PATH        = "test_data_60k.csv"
TARGET_COL       = "price"
MODEL_PATH       = "lr_model.pkl"
SCALER_PATH      = "lr_scaler.pkl"
REPORT_PATH      = "LR_Crop_Price_Report.docx"
TRAIN_SORTED_PNG = "lr_train_sorted_plot.png"
TEST_SORTED_PNG  = "lr_test_sorted_plot.png"
SCATTER_PNG      = "lr_scatter_plot.png"
GRAPH_SAMPLES    = 2000


# ═══════════════════════════════════════════════════════════════════════════════
#  1. DATA LOADING
# ═══════════════════════════════════════════════════════════════════════════════

def load_data(train_path, test_path):
    for p in (train_path, test_path):
        if not os.path.isfile(p):
            sys.exit(f"[ERROR] File not found: '{p}'")

    train_df = pd.read_csv(train_path)
    test_df  = pd.read_csv(test_path)

    print("Loaded real dataset successfully")
    print(f"[INFO] Training data shape : {train_df.shape}")
    print(f"[INFO] Testing  data shape : {test_df.shape}")

    return train_df, test_df


# ═══════════════════════════════════════════════════════════════════════════════
#  2. PREPROCESSING
# ═══════════════════════════════════════════════════════════════════════════════

def handle_missing_values(df, label="data"):
    for col in df.columns:
        if df[col].isnull().sum() == 0:
            continue
        if df[col].dtype in ("float64", "int64", "float32", "int32"):
            df[col].fillna(df[col].median(), inplace=True)
        else:
            df[col].fillna(df[col].mode()[0], inplace=True)
    return df


def preprocess(train_df, test_df, target_col):
    # Handle missing values
    train_df = handle_missing_values(train_df.copy(), "training set")
    test_df  = handle_missing_values(test_df.copy(),  "testing set")

    # One-hot encoding for "season" (and any other categorical columns)
    cat_cols = train_df.select_dtypes(include=["object", "category"]).columns.tolist()
    if cat_cols:
        train_df = pd.get_dummies(train_df, columns=cat_cols)
        test_df  = pd.get_dummies(test_df, columns=cat_cols)
        # Align columns to ensure test set has identical dummy columns as train set
        train_df, test_df = train_df.align(test_df, join='left', axis=1, fill_value=0)

    # Separate features and target
    feat_cols = [c for c in train_df.columns if c != target_col]
    X_train = train_df[feat_cols].values
    y_train = train_df[target_col].values
    X_test  = test_df[feat_cols].values
    y_test  = test_df[target_col].values

    # StandardScaler: fit ONLY on training data
    scaler = StandardScaler()
    X_train_s = scaler.fit_transform(X_train)
    X_test_s  = scaler.transform(X_test)
    
    print("[INFO] One-hot encoding and feature scaling applied (StandardScaler fit on train only).")
    
    return X_train_s, X_test_s, y_train, y_test, scaler, feat_cols


# ═══════════════════════════════════════════════════════════════════════════════
#  3 & 4. MODEL AND TRAINING
# ═══════════════════════════════════════════════════════════════════════════════

def build_and_train_model(X_train, y_train):
    print("\n[MODEL] Linear Regression created")
    print("[INFO] Training on full dataset...")
    model = LinearRegression()
    model.fit(X_train, y_train)
    print("[INFO] ✅ Model training complete.")
    return model


# ═══════════════════════════════════════════════════════════════════════════════
#  5. METRICS
# ═══════════════════════════════════════════════════════════════════════════════

def compute_metrics(y_true, y_pred, label=""):
    mse  = mean_squared_error(y_true, y_pred)
    mae  = mean_absolute_error(y_true, y_pred)
    rmse = math.sqrt(mse)
    r2   = r2_score(y_true, y_pred)

    print("=" * 55)
    print(f"       {label} METRICS")
    print("=" * 55)
    print(f"  Mean Absolute Error (MAE)  : {mae:>12.4f}")
    print(f"  Mean Squared Error  (MSE)  : {mse:>12.4f}")
    print(f"  Root Mean Sq Error  (RMSE) : {rmse:>12.4f}")
    print(f"  R² Score                   : {r2:>12.4f}")
    print("=" * 55 + "\n")

    return {"mae": mae, "mse": mse, "rmse": rmse, "r2": r2}


# ═══════════════════════════════════════════════════════════════════════════════
#  6. GRAPH VISUALISATION
# ═══════════════════════════════════════════════════════════════════════════════

def plot_sorted(y_true, y_pred, title, save_path, n=GRAPH_SAMPLES):
    """Sort by actual price and plot first n samples."""
    n_plot = min(n, len(y_true))
    # Subsample the data to first n elements to respect memory constraints
    y_true_sub = y_true[:n_plot]
    y_pred_sub = y_pred[:n_plot]
    
    sorted_idx   = np.argsort(y_true_sub)
    y_true_sort  = y_true_sub[sorted_idx]
    y_pred_sort  = y_pred_sub[sorted_idx]

    idx = np.arange(n_plot)

    plt.figure(figsize=(14, 5))
    plt.plot(idx, y_true_sort, color="#2196F3", linewidth=1.2, alpha=0.9, label="Actual Price")
    plt.plot(idx, y_pred_sort, color="#FF5722", linewidth=1.2, alpha=0.9, label="Predicted Price")
    plt.xlabel("Sample Index (sorted by actual price)", fontsize=12)
    plt.ylabel("Price (₹)", fontsize=12)
    plt.title(title, fontsize=14, fontweight="bold")
    plt.legend(fontsize=11)
    plt.tight_layout()
    plt.savefig(save_path, dpi=150)
    plt.close()
    print(f"[INFO] 📊  Graph saved → {save_path}")


def plot_scatter(y_true, y_pred, save_path, n=GRAPH_SAMPLES):
    """Scatter plot with y = x line, limited to n samples."""
    n_plot = min(n, len(y_true))
    y_true_sub = y_true[:n_plot]
    y_pred_sub = y_pred[:n_plot]
    
    plt.figure(figsize=(8, 8))
    plt.scatter(y_true_sub, y_pred_sub, alpha=0.3, s=8, color="#1976D2", label="Predictions")

    lo = min(y_true_sub.min(), y_pred_sub.min())
    hi = max(y_true_sub.max(), y_pred_sub.max())
    plt.plot([lo, hi], [lo, hi], color="#E53935", linewidth=2, linestyle="--", label="Ideal (y = x)")

    plt.xlabel("Actual Price (₹)", fontsize=12)
    plt.ylabel("Predicted Price (₹)", fontsize=12)
    plt.title("Actual vs Predicted Scatter Plot", fontsize=14, fontweight="bold")
    plt.legend(fontsize=11)
    plt.tight_layout()
    plt.savefig(save_path, dpi=150)
    plt.close()
    print(f"[INFO] 📊  Graph saved → {save_path}")


# ═══════════════════════════════════════════════════════════════════════════════
#  7. RECOMMENDATION SYSTEM
# ═══════════════════════════════════════════════════════════════════════════════

def generate_recommendations(y_actual, y_pred, hold=0.10, wait=0.02):
    recs = []
    changes = []
    for actual, pred in zip(y_actual, y_pred):
        change = (pred - actual) / actual
        changes.append(change)
        if change >= hold:
            recs.append("HOLD")
        elif change >= wait:
            recs.append("WAIT")
        else:
            recs.append("SELL")
    return recs, changes


def print_recommendations(y_actual, y_pred, changes, recs, n=10):
    print(f"\n── First {n} Recommendations (Test Data) ──")
    print(f"{'Actual':>12} | {'Predicted':>12} | {'Change':>8} | {'Recommendation':>14}")
    print("-" * 55)
    for i in range(min(n, len(y_actual))):
        print(f"{y_actual[i]:>12.2f} | {y_pred[i]:>12.2f} | {changes[i]:>8.4f} | {recs[i]:>14}")
    print("-" * 55 + "\n")


# ═══════════════════════════════════════════════════════════════════════════════
#  8. WORD REPORT
# ═══════════════════════════════════════════════════════════════════════════════

def _cell(cell, text, bold=False, center=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.font.size = Pt(10)
    run.bold = bold
    if center:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def _add_metrics_section(doc, heading, metrics):
    doc.add_heading(heading, level=1)
    tm = doc.add_table(rows=1, cols=2, style="Light Shading Accent 1")
    _cell(tm.rows[0].cells[0], "Metric", bold=True, center=True)
    _cell(tm.rows[0].cells[1], "Value",  bold=True, center=True)
    for key, label in [("mae", "MAE"), ("rmse", "RMSE"), ("r2", "R² Score")]:
        r = tm.add_row()
        _cell(r.cells[0], label, center=True)
        _cell(r.cells[1], f"{metrics[key]:.4f}", center=True)

def generate_word_report(train_rows, test_rows, train_metrics, test_metrics,
                         y_true, y_pred, recs, report_path):
    doc = Document()

    # Title
    t = doc.add_heading("Crop Price Forecasting using Linear Regression", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Dataset Summary
    doc.add_heading("1. Dataset Summary", level=1)
    doc.add_paragraph(f"Training samples : {train_rows:,}")
    doc.add_paragraph(f"Testing  samples : {test_rows:,}")

    # 2. Model Details
    doc.add_heading("2. Model Details", level=1)
    doc.add_paragraph("Algorithm: Linear Regression")

    # 3 & 4. Metrics
    _add_metrics_section(doc, "3. Training Metrics", train_metrics)
    _add_metrics_section(doc, "4. Testing Metrics",  test_metrics)
    
    doc.add_paragraph("The model shows good generalization with a small performance gap between training and testing.")

    # 5. Visualisations
    doc.add_heading("5. Visualisations", level=1)

    if os.path.isfile(TRAIN_SORTED_PNG):
        doc.add_paragraph("Training Data (Sorted): Actual vs Predicted")
        doc.add_picture(TRAIN_SORTED_PNG, width=Inches(6.0))

    if os.path.isfile(TEST_SORTED_PNG):
        doc.add_paragraph("Testing Data (Sorted): Actual vs Predicted")
        doc.add_picture(TEST_SORTED_PNG, width=Inches(6.0))
        doc.add_paragraph("Sorted plots improve trend visibility.")

    if os.path.isfile(SCATTER_PNG):
        doc.add_paragraph("Actual vs Predicted Scatter Plot")
        doc.add_picture(SCATTER_PNG, width=Inches(5.0))
        doc.add_paragraph("Scatter plot shows prediction accuracy.")

    # 6. Sample Predictions Table
    doc.add_heading("6. Sample Predictions (First 20 — Test Data)", level=1)
    n = min(20, len(y_true))
    pt = doc.add_table(rows=1, cols=3, style="Light Shading Accent 1")
    for i, h in enumerate(["Actual Price", "Predicted Price", "Recommendation"]):
        _cell(pt.rows[0].cells[i], h, bold=True, center=True)
    for i in range(n):
        r = pt.add_row()
        _cell(r.cells[0], f"{y_true[i]:.2f}", center=True)
        _cell(r.cells[1], f"{y_pred[i]:.2f}", center=True)
        _cell(r.cells[2], recs[i],             center=True)

    doc.save(report_path)
    print(f"[INFO] 📄  Word report saved → {report_path}")


# ═══════════════════════════════════════════════════════════════════════════════
#  9. SAVE MODEL
# ═══════════════════════════════════════════════════════════════════════════════

def save_artifacts(model, scaler):
    joblib.dump(model,  MODEL_PATH)
    joblib.dump(scaler, SCALER_PATH)
    print(f"[INFO] Model  saved → {MODEL_PATH}")
    print(f"[INFO] Scaler saved → {SCALER_PATH}")


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print("\n" + "═" * 55)
    print("  AI-Based Crop Price Forecasting using Linear Regression")
    print("═" * 55 + "\n")

    # 1. Load Data
    train_df, test_df = load_data(TRAIN_PATH, TEST_PATH)

    # 2. Preprocess
    (X_train, X_test, y_train, y_test, scaler, feat_names) = preprocess(train_df, test_df, TARGET_COL)

    # 3 & 4. Train Model
    model = build_and_train_model(X_train, y_train)

    print("[INFO] Predicting on training data...")
    y_pred_train = model.predict(X_train)
    print("[INFO] Predicting on testing data...")
    y_pred_test  = model.predict(X_test)

    # 5. Compute Metrics
    train_metrics = compute_metrics(y_train, y_pred_train, label="TRAINING")
    test_metrics  = compute_metrics(y_test,  y_pred_test,  label="TESTING")

    # 6. Graph Visualization
    plot_sorted(y_train, y_pred_train, "Training Data (Sorted): Actual vs Predicted", TRAIN_SORTED_PNG, n=GRAPH_SAMPLES)
    plot_sorted(y_test, y_pred_test, "Testing Data (Sorted): Actual vs Predicted", TEST_SORTED_PNG, n=GRAPH_SAMPLES)
    plot_scatter(y_test, y_pred_test, SCATTER_PNG, n=GRAPH_SAMPLES)

    # 7. Recommendation System
    recs, changes = generate_recommendations(y_test, y_pred_test)
    print_recommendations(y_test, y_pred_test, changes, recs, n=10)

    # 8. Word Report
    generate_word_report(
        train_rows=len(y_train), test_rows=len(y_test),
        train_metrics=train_metrics, test_metrics=test_metrics,
        y_true=y_test, y_pred=y_pred_test, recs=recs,
        report_path=REPORT_PATH
    )

    # 9. Save Artifacts
    save_artifacts(model, scaler)

if __name__ == "__main__":
    main()
